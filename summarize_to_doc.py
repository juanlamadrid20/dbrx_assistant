import json
import os
import requests
import sys
from dotenv import load_dotenv
from docx import Document

# Load environment variables
load_dotenv()

# Configuration
DATABRICKS_TOKEN = os.getenv("DATABRICKS_TOKEN")
DATABRICKS_URL = os.getenv("DATABRICKS_URL")
NOTES_DIR = os.getenv("NOTES_DIR")
SUMMARY_DOC_DIR = os.getenv("SUMMARY_DOC_DIR")
HEADERS = {
    "Authorization": f"Bearer {DATABRICKS_TOKEN}",
    "Content-Type": "application/json",
}


def read_transcribe_file(trans_file):
    """Read content from a file."""
    with open(trans_file, "r") as file:
        return file.read()


def request_to_databricks(data):
    """Send a formatted request to the Databricks API."""
    response = requests.post(DATABRICKS_URL, headers=HEADERS, data=json.dumps(data, allow_nan=True))
    if response.status_code != 200:
        raise Exception(f"Request failed with status {response.status_code}: {response.text}")
    return response.json()["choices"][0]["message"]["content"]


def perform_extraction(transcription, role_description, max_tokens=4096):
    """Extract information based on the specified role and transcription."""
    data = {
        "messages": [
            {"role": "system", "content": role_description},
            {"role": "user", "content": transcription}
        ],
        "temperature": 0.5,
        "top_p": 0.95,
        "max_tokens": max_tokens
    }
    return request_to_databricks(data)


def extract_information(transcription):
    """Extract various types of information from transcription."""
    responses = {
        "abstract_summary": perform_extraction(
            transcription,
            "You are a highly skilled AI trained in language comprehension and summarization. Please read the following text and summarize it into a concise abstract paragraph. Avoid unnecessary details."
        ),
        "key_points": perform_extraction(
            transcription,
            "You are a databricks AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points discussed."
        ),
        "action_items": perform_extraction(
            transcription,
            "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon."
        ),
        "sentiment": perform_extraction(
            transcription,
            "Analyze the sentiment of the following text, considering the overall tone and the emotion conveyed by the language used."
        ),
        "email": perform_extraction(
            transcription,
            "Write a follow-up email summarizing the text."
        )
    }
    return responses


def save_document(data, filename, format='docx'):
    """Save extracted data into a document, either DOCX or text."""
    if format == 'docx':
        doc = Document()
        for key, value in data.items():
            heading = key.replace("_", " ").title()
            doc.add_heading(heading, level=1)
            doc.add_paragraph(value)
            doc.add_paragraph()
        doc.save(filename)
    else:
        with open(filename, 'w') as file:
            for key, value in data.items():
                heading = key.replace("_", " ").title()
                file.write(f"{heading}\n{value}\n\n")


def run_me(filename):
    """Main function to run the whole process."""
    file_path = os.path.join(NOTES_DIR, filename)
    transcription = read_transcribe_file(file_path)
    data = extract_information(transcription)

    base_filename = os.path.splitext(filename)[0]
    docx_filename = os.path.join(SUMMARY_DOC_DIR, f"{base_filename}_output.docx")
    txt_filename = os.path.join(SUMMARY_DOC_DIR, f"{base_filename}_output.txt")

    save_document(data, docx_filename)
    save_document(data, txt_filename, format='txt')


if __name__ == "__main__":
    run_me(sys.argv[1])
